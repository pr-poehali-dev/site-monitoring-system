#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import requests
from docx import Document
from io import BytesIO
import json

def handler(request):
    """Анализирует DOCX файл и ищет упоминания документов"""
    
    url = "https://cdn.poehali.dev/projects/43869288-534e-4e87-af2a-4115997b9a30/bucket/docs/Постановления/926_main_4ae01c6e.docx"
    
    result = {}
    try:
        # Скачиваем файл
        response = requests.get(url, timeout=60)
        response.raise_for_status()
        result['file_size'] = len(response.content)
        docx_file = BytesIO(response.content)
        
        # Извлекаем первые 20 параграфов
        doc = Document(docx_file)
        result['total_paragraphs'] = len(doc.paragraphs)
        paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            if len(paragraphs) >= 20:
                break
            if para.text.strip():
                paragraphs.append({
                    'index': i,
                    'text': para.text.strip()
                })
        
        # Объединяем текст
        full_text = '\n'.join([p['text'] for p in paragraphs])
        
        # Ищем упоминания по трем паттернам
        mentions = []
        
        # Паттерн 1
        pattern1 = r'от\s+(\d{2}\.\d{2}\.\d{4})\s+года?\s+(?:№|N|#)\s*(\d+)'
        for match in re.finditer(pattern1, full_text, re.IGNORECASE | re.UNICODE):
            mentions.append({
                'pattern': 'Паттерн 1 (обратный)',
                'date': match.group(1),
                'number': match.group(2),
                'full_match': match.group(0)
            })
        
        # Паттерн 2
        pattern2 = r'постановлени[ея]\s+(?:№|N|#)?\s*(\d+)\s+от\s+(\d{2}\.\d{2}\.\d{4})'
        for match in re.finditer(pattern2, full_text, re.IGNORECASE | re.UNICODE):
            mentions.append({
                'pattern': 'Паттерн 2 (прямой)',
                'number': match.group(1),
                'date': match.group(2),
                'full_match': match.group(0)
            })
        
        # Паттерн 3
        pattern3 = r'(?:№|N|#)\s*(\d+)\s+от\s+(\d{2}\.\d{2}\.\d{4})'
        for match in re.finditer(pattern3, full_text, re.IGNORECASE | re.UNICODE):
            mentions.append({
                'pattern': 'Паттерн 3 (упрощенный)',
                'number': match.group(1),
                'date': match.group(2),
                'full_match': match.group(0)
            })
        
        # Анализ формата
        analysis = {
            'has_numbers': bool(re.search(r'\d+', full_text)),
            'has_dates': bool(re.search(r'\d{2}\.\d{2}\.\d{4}', full_text)),
            'has_number_signs': bool(re.search(r'[№N#]', full_text)),
            'has_postanovlenie': bool(re.search(r'постановлени[ея]', full_text, re.IGNORECASE)),
            'has_ot': bool(re.search(r'\bот\b', full_text, re.IGNORECASE))
        }
        
        result.update({
            'paragraphs': paragraphs,
            'mentions': mentions,
            'analysis': analysis,
            'full_text': full_text[:5000]  # Ограничиваем размер
        })
        
        return {
            'statusCode': 200,
            'body': json.dumps(result, ensure_ascii=False)
        }
        
    except Exception as e:
        import traceback
        return {
            'statusCode': 500,
            'body': json.dumps({
                'error': str(e),
                'traceback': traceback.format_exc()
            }, ensure_ascii=False)
        }